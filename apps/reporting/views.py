from collections import defaultdict
import io
import json
import os
import pickle
import re
import copy
# For file conversion
import subprocess # Libreoffice
from docx2pdf import convert # MS Word

from tempfile import NamedTemporaryFile
#from tkinter import Image
from django.http import FileResponse, HttpResponse, HttpResponseRedirect
from django.shortcuts import render, get_object_or_404
import datetime
import docx
from docx import Document
from docx.shared import Inches
from django.core.exceptions import ValidationError
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from django.views.generic import ListView, DetailView, CreateView
import numpy as np
import pandas as pd
from sklearn.tree import export_graphviz
from shapely.geometry import Polygon, Point

#from SOM.utils import get_uicompo_from_centroid
from apps.decisiondiscovery.utils import truncar_a_dos_decimales
from apps.notification.views import create_notification
from core.utils import read_ui_log_as_dataframe
from core.settings import PROCESS_DISCOVERY_LOG_FILENAME
from .models import PDD
from django.utils.translation import gettext_lazy as _
from apps.analyzer.models import Execution 
from .forms import ReportingForm
from django.urls import reverse

import os
import zipfile
from django.http import HttpResponse
from django.core.exceptions import ValidationError
from django.utils.translation import gettext as _
from django.conf import settings
from django.shortcuts import get_object_or_404

import pydotplus
#import pypandoc
from PIL import Image, ImageDraw
#import subprocess
import aspose.words as aw
##################################
from graphviz import Source
from tempfile import NamedTemporaryFile
from apps.processdiscovery.utils import Process, extract_all_activities_labels, extract_prev_act_labels

import pygraphviz as pgv 
from docx.shared import RGBColor

#########################################
########## COMPONENT RETRIEVAL ##########
#########################################

def get_som_json_from_screenshots_in_components(screenshot_filename, scenario_results_path) -> dict:
    
    if not os.path.exists(os.path.join(scenario_results_path, "components_json")):
        raise Exception("No UI Elm. Det. Phase Conducted")

    return json.load(open(os.path.join(scenario_results_path, "components_json", f"{screenshot_filename}.json")))


def get_uicompo_from_centroid(screenshot_filename, ui_compo_centroid, ui_compo_class_or_text, scenario_results_path) -> dict:
    ui_compo_centroid = [int(float(coord)) for coord in ui_compo_centroid]
    som_json = get_som_json_from_screenshots_in_components(screenshot_filename, scenario_results_path)

    uicompo_json = None
    min_distance = float('inf')
    closest_compo = None

    class_compos = list(filter(lambda compo: compo['class'] == ui_compo_class_or_text, som_json['compos']))
    # If none is found, that means the class is actually a text, or NaN
    if len(class_compos) == 0 and ui_compo_class_or_text not in ["NaN", "nan"]:
        class_compos = list(filter(lambda compo: compo['text'] == ui_compo_class_or_text, som_json['compos']))
    # It could happen that in this specific instance, the object was not detected. In this case we will just get the closest one
    if len(class_compos) == 0:
        class_compos = som_json['compos']

    for compo in class_compos:
        # Convertir el centroid de compo y ui_compo_centroid a enteros antes de comparar
        compo_centroid_int = [int(float((coord))) for coord in compo['centroid']]
        if compo_centroid_int == ui_compo_centroid:
            uicompo_json = compo
            break
        else:
            # Calcular la distancia entre centroids
            distance = np.linalg.norm(np.array(compo_centroid_int) - np.array(ui_compo_centroid))
            if distance < min_distance:
                # Check the centroid is inside the bounding box
                compo_polygon = Polygon(compo['points'])
                if compo_polygon.contains(Point(ui_compo_centroid)):
                    min_distance = distance
                    closest_compo = compo

    # Si no se encontró una coincidencia exacta, usar el más cercano
    if uicompo_json is None and closest_compo is not None:
        uicompo_json = closest_compo
    return uicompo_json

def pre_pd_activities(decision_point, df_pd_log,colnames):
    grouped = df_pd_log.groupby(colnames['Variant'])
    # Iterate through each group
    result = []
    for variant, group in grouped:
        # Sort by Timestamp
        sorted_group = group.sort_values(by=colnames['Timestamp'])
        # Iterate through activities in the sorted group
        for activity in sorted_group[colnames['Activity']]:
            
            if activity not in result:
                result.append(activity)
            if str(activity) == decision_point:
                break
            # Stop when encountering the specific activity
            
    return result

def pdd_define_style(document):
    # =======================================================================================================
    # add a custom style for headings
    heading_style = document.styles.add_style('Heading', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.base_style = document.styles['Heading 1']
    heading_style.font.name = 'Calibri Light'
    heading_style.font.size = docx.shared.Pt(18)
    heading_style.font.color.rgb = docx.shared.RGBColor(0x2d, 0x89, 0xff)

    # add a custom style for the table of contents heading
    toc_heading_style = document.styles.add_style('indice', WD_STYLE_TYPE.PARAGRAPH)
    toc_heading_style.base_style = document.styles['Heading']
    toc_heading_style.font.color.rgb = docx.shared.RGBColor(0x2d, 0x89, 0xff)
    toc_heading_style.font.bold = True
    toc_heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # add a custom style for section headings
    section_heading_style = document.styles.add_style('Section Heading', WD_STYLE_TYPE.PARAGRAPH)
    section_heading_style.base_style = document.styles['Heading 2']
    section_heading_style.font.name = 'Calibri Light'
    section_heading_style.font.size = docx.shared.Pt(14)
    section_heading_style.font.color.rgb = docx.shared.RGBColor(0x2d, 0x89, 0xff)
    section_heading_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    section_heading_style.paragraph_format.line_spacing = 1.5

    # add a custom style for body text
    body_style = document.styles.add_style('normal', WD_STYLE_TYPE.PARAGRAPH)
    body_style.base_style = document.styles['Normal']
    body_style.font.name = 'Calibri'
    body_style.font.size = docx.shared.Pt(12)
    # =======================================================================================================
    
    # Create a custom document style
    doc_styles = document.styles
    custom_style = doc_styles.add_style('CustomStyle', 1)
    custom_style.font.name = 'Segoe UI'
    custom_style.font.size = 12
    
    return document, custom_style

def pdd_define_properties(document):
    # add document properties
    document.core_properties.author = "Antonio Martinez Rojas"
    document.core_properties.title = _("Process Definition Document")
    document.core_properties.comments = _("This is a first version of your RPA Process Definition Document.")
    document.core_properties.category = _("PDDs")
    document.core_properties.created = datetime.datetime.now()
    
    return document
    
def pdd_add_cover(document, custom_style):
    
    # Set the page orientation to portrait
    section = document.sections[0]
    section.orientation = 1
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    
    # Add a cover page
    document.add_picture('cover_image.png', width=Inches(6))
    document.add_paragraph(_('Process Definition Document'), custom_style).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(_('Version 1.0'), custom_style).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()
    
    return document
    
def pdd_table_of_contents(document):
    # add a table of contents
    paragraph = document.add_paragraph()
    run = paragraph.add_run(_('Table of Contents'))
    paragraph.style = 'indice'
    # paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Inches(0.2)
    paragraph.paragraph_format.space_before = Inches(0.2)
    # Code for making Table of Contents
    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:t')
    fldChar3.text = _("Right-click to update field.")

    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

    p_element = paragraph._p
    document.add_page_break()
    
    return document

def ui_screen_trace_back_reporting(case_study_id):
    # create a new document
    document = Document()
    
    document = pdd_define_properties(document)
    document, custom_style = pdd_define_style(document)
    document = pdd_add_cover(document, custom_style)
    document = pdd_table_of_contents(document)
    
    # SECTIONS

    # add section 1: introduction
    document.add_heading(_('1. Introduction'), level=1).style = 'Section Heading'
    # section_heading = document.add_paragraph()
    # section_heading.style = 'Section Heading'
    # section_heading.add_run('1. Introduction')

    # add body text for section 1
    body_text = document.add_paragraph()
    body_text.style = 'Body Text'
    body_text.add_run(_('This document outlines the process definition for...'))
    document.add_page_break()

    # add section 2: process overview
    document.add_heading('2. Process Overview', level=1).style = 'Section Heading'
    # section_heading = document.add_paragraph()
    # section_heading.style = 'Section Heading'
    # section_heading.add_run('2. Process Overview')

    # add body text for section 2
    body_text = document.add_paragraph()
    body_text.style = 'normal'
    body_text.add_run(_('This process consists of the following steps:'))
    # add a numbered list for the steps in section 2
    numbered_list = document.add_paragraph()
    numbered_list.style = 'List Number'
    numbered_list.add_run(_('Step 1: ...'))
    numbered_list.add_run(_('\nStep 2: ...'))
    numbered_list.add_run(_('\nStep 3: ...'))
    
    # add section 3: roles and responsibilities
    document.add_page_break()
    document.add_heading(_(), level=1).style = 'Section Heading'
    # section_heading = document.add_paragraph()
    # section_heading.style = 'Section Heading'
    # section_heading.add_run('3. Roles and Responsibilities')

    # add body text for section 3
    body_text = document.add_paragraph()
    body_text.style = 'Body Text'
    body_text.add_run(_('The following roles and responsibilities are involved in this process:'))

    # add a table for the roles and responsibilities in section 3
    table = document.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.autofit = False

    # set the width of the columns in the table
    column_widths = (Inches(1.5), Inches(4))
    for i, width in enumerate(column_widths):
        table.columns[i].width = width

    # add the header row for the table
    heading_cells = table.rows[0].cells
    heading_cells[0].text = _('Role')
    heading_cells[1].text = _('Responsibilities')

    # add the content for the remaining rows in the table
    row_cells = table.rows[1].cells
    row_cells[0].text = _('Manager')
    row_cells[1].text = _('Approves final output of the process')
    row_cells = table.rows[2].cells
    row_cells[0].text = _('Analyst')
    row_cells[1].text = _('Conducts analysis for the process')
    row_cells = table.rows[3].cells
    row_cells[0].text = _('Developer')
    row_cells[1].text = _('Implements changes to the process')

    # add section 4: process flow
    document.add_page_break()
    document.add_heading(_('4. Process Flow'), level=1).style = 'Section Heading'
    # section_heading = document.add_paragraph()
    # section_heading.style = 'Section Heading'
    # section_heading.add_run('4. Process Flow')

    # add body text for section 4
    body_text = document.add_paragraph()
    body_text.style = 'Body Text'
    body_text.add_run(_('The process flow for this process is as follows:'))
    # add an image for the process flow in section 4
    document.add_picture('process_flow.png', width=Inches(6))
    # add section 5: process metrics
    document.add_page_break()
    document.add_heading(_('5. Process Metrics'), level=1).style = 'Section Heading'
    # section_heading = document.add_paragraph()
    # section_heading.style = 'Section Heading'
    # section_heading.add_run('5. Process Metrics')

    # add body text for section 5
    body_text = document.add_paragraph()
    body_text.style = 'Body Text'
    body_text.add_run(_('The following metrics are tracked for this process:'))
    # add a bulleted list for the metrics in section 5
    bulleted_list = document.add_paragraph()
    bulleted_list.style = 'List Bullet'
    # add the metrics to the bulleted list
    bulleted_list.add_run(_('Metric 1: ...'))
    bulleted_list.add_run(_('\nMetric 2: ...'))
    bulleted_list.add_run(_('\nMetric 3: ...'))
    
    # add section 6: process improvement
    document.add_page_break()
    document.add_heading(_('6. Process Improvement'), level=1).style = 'Section Heading'
    # section_heading = document.add_paragraph()
    # section_heading.style = 'Section Heading'
    # section_heading.add_run('6. Process Improvement')

    # add body text for section 6
    body_text = document.add_paragraph()
    body_text.style = 'Body Text'
    body_text.add_run(_('The following improvement opportunities have been identified for this process:'))
    # add a numbered list for the improvement opportunities in section 6

    # Adding list of style name 'List Number'
    document.add_heading('Style: List Number', 3)
    # Adding points to the list named 'List Number'
    document.add_paragraph(_('Opportunity 1: ...'), style='List Number')
    document.add_paragraph(_('Opportunity 2: ...'), style='List Number')
    document.add_paragraph(_('Opportunity 3: ...'), style='List Number')
    
    # Adding list of style name 'List Number 2'
    document.add_heading('Style: List Number 2', 3)
    # Adding points to the list named 'List Number 2'
    document.add_paragraph(_('Opportunity 1: ...'), style='List Number 2')
    document.add_paragraph(_('Opportunity 2: ...'), style='List Number 2')
    document.add_paragraph(_('Opportunity 3: ...'), style='List Number 2')
    
    # Adding list of style name 'List Number 3'
    document.add_heading('Style: List Number 3', 3)
    # Adding points to the list named 'List Number 3'
    document.add_paragraph(_('Opportunity 1: ...'), style='List Number 3')
    document.add_paragraph(_('Opportunity 2: ...'), style='List Number 3')
    document.add_paragraph(_('Opportunity 3: ...'), style='List Number 3')    


    # add section 7: glossary
    document.add_page_break()
    document.add_heading(_(), level=1).style = 'Section Heading'
    # section_heading = document.add_paragraph()
    # section_heading.style = 'Section Heading'
    # section_heading.add_run('7. Glossary')
    
    # add body text for section 7
    body_text = document.add_paragraph()
    body_text.style = 'Body Text'
    body_text.add_run(_('The following terms are used throughout this document:'))
    
    # add a table for the glossary in section 7
    table = document.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    # set the width of the columns in the table

    column_widths = (Inches(1.5), Inches(4))

    for i, width in enumerate(column_widths):
            table.columns[i].width = width
            
    # add the header row for the table
    heading_cells = table.rows[0].cells
    heading_cells[0].text = _('Term')
    heading_cells[1].text = _('Definition')
    
    # add the content for the remaining rows in the table
    row_cells = table.rows[1].cells
    row_cells[0].text = _('Process')
    row_cells[1].text = _('A set of activities that transform inputs into outputs')
    row_cells = table.rows[2].cells
    row_cells[0].text = _('Metric')
    row_cells[1].text = _('A quantifiable measure used to track the performance of a process')

    # save the document
    document.save('Process Definition Document.docx')
    
    return document
    
    
# TODO:
class ReportGenerateView(DetailView):
    def get(self, request, *args, **kwargs):
        cs = get_object_or_404(PDD, id=kwargs["case_study_id"], active=True)
        
        document = ui_screen_trace_back_reporting(cs.id)
        
        # check if the file exists
        if not os.path.exists(document.file.path):
            return HttpResponse("File does not exist", status=404)

        # set the file content type and headers
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{document.name}"'
        
        # read the file and write it to the response
        with open(document.file.path, 'rb') as file:
            response.write(file.read())
        
        return response

# TODO:
class ReportDownloadView(DetailView):
    def get(self, request, *args, **kwargs):
        pdd = get_object_or_404(PDD, id=kwargs["report_id"], active=True)
        context = {"pdd": pdd}
        
        # check if the file exists
        if not os.path.exists(pdd.pdd.file.path):
            return HttpResponse(_("File does not exist"), status=404)

        # set the file content type and headers
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{pdd.pdd.name}"'
        
        # read the file and write it to the response
        with open(pdd.pdd.file.path, 'rb') as file:
            response.write(file.read())
        
        return response

# TODO:
class ReportListView(ListView):
    model = PDD
    template_name = "reporting/list.html"
    paginate_by = 50

    def get_context_data(self, **kwargs):
        context = super(ReportListView, self).get_context_data(**kwargs)
        context['case_study_id'] = self.kwargs.get('case_study_id')
        return context

    def get_queryset(self):
        # Obtiene el ID del Experiment pasado como parámetro en la URL
        case_study_id = self.kwargs.get('case_study_id')

        # Search if s is a query parameter
        search = self.request.GET.get("s")
        # Filtra los objetos por case_study_id
        if search:
            queryset = PDD.objects.filter(case_study__id=case_study_id, case_study__user=self.request.user, case_study__title__icontains=search).order_by('-created_at')
        else:
            queryset = PDD.objects.filter(case_study__id=case_study_id, case_study__user=self.request.user).order_by('-created_at')
        
        return queryset

#############################################################################################

def tree_to_png(path_to_tree_file):
    # Load the decision tree data
    try:
        with open('/screenrpa/' + path_to_tree_file, 'rb') as file:
            loaded_data = pickle.load(file)
        loaded_classifier = loaded_data['classifier']
        loaded_feature_names = loaded_data['feature_names']
        loaded_class_names = [str(item) for item in loaded_data['class_names']]
    except FileNotFoundError:
        print(f"File not found: {path_to_tree_file}")
        return None
    
    dot_data = io.StringIO()
    export_graphviz(loaded_classifier, out_file=dot_data, filled=True, rounded=True,
                    special_characters=True, feature_names=loaded_feature_names, class_names=loaded_class_names)
    graph = pydotplus.graph_from_dot_data(dot_data.getvalue())
    png_image = graph.create_png()

    # Save the image to a temporary file
    temp_file = NamedTemporaryFile(delete=False, suffix='.png')
    with open(temp_file.name, 'wb') as file:
        file.write(png_image)
    
    return temp_file.name


#########################################
def dot_to_png(dot_path):
    # Cargar el contenido del archivo .dot
    with open(dot_path, 'r') as file:
        dot_content = file.read()
        try:
            graph = Source(dot_content)
            graph.format = 'png'
            
            # Guardar la imagen a un archivo temporal
            temp_file = NamedTemporaryFile(delete=False, suffix='.png')
            temp_file.close()
            graph_path = graph.render(filename=temp_file.name, format='png', cleanup=True)
            
            return graph_path 
        
        except Exception as e:

            print(f"Error al procesar el gráfico: {e}")
            return None

#############################################################################################
class ReportCreateView(CreateView):
    model = PDD
    form_class = ReportingForm
    template_name = "reporting/create.html"

    def get_context_data(self, **kwargs):
        context = super(ReportCreateView, self).get_context_data(**kwargs)
        
        execution_id = self.kwargs.get('execution_id')
        execution = Execution.objects.get(pk=execution_id)
        context['execution'] = execution
        reports = PDD.objects.filter(execution=execution).order_by('-created_at')
        context['reports'] = reports
        return context 
    
    def form_valid(self, form):
        if not self.request.user.is_authenticated:
            raise ValidationError(_("User must be authenticated."))
        self.object = form.save(commit=False)
        self.object.user = self.request.user
        self.object.execution = Execution.objects.get(pk=self.kwargs.get('execution_id'))

        saved = self.object.save()

        self.report_generate(self.object)
        
        return HttpResponseRedirect(self.get_success_url())
    
    def report_generate(self, report):
        execution= self.get_context_data()['execution']
        scenarios_to_study= self.get_context_data()['execution'].scenarios_to_study

        for scenario in scenarios_to_study:

            report_directory = os.path.join(execution.exp_folder_complete_path, scenario+"_results")
            report_path = os.path.join(report_directory, f'report_{report.id}.docx')

            # Ensure the directory exists
            os.makedirs(report_directory, exist_ok=True)

            report_define(report_directory, report_path, execution, report, scenario)
            # with open(report_path, 'wb') as file:
            #     file.write(b'PDF content or whatever content you generate')
        

def report_define(report_directory, report_path, execution,  report, scenario):
    template_path = os.path.join("apps", "templates", "reporting", "report_template.docx")
    doc = Document(template_path)
    colnames=execution.case_study.special_colnames
    ###############3
    paragraph_dict = {}
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        if text.startswith('[') and text.endswith(']'): 
            paragraph_dict[text] = i 

    for key, value in paragraph_dict.items():
        print(f"Texto del párrafo: {key}\nÍndice: {value}\n")

    ############################ INTRODUCION
    
    purpose= doc.paragraphs[paragraph_dict['[PURPOSE]']]
    purpose.text = report.purpose
    for style in doc.styles:
        print(f"Style name: {style.name}, Style type: {style.type}")
    
    #paragraph.style = doc.styles['Normal']

    purpose= doc.paragraphs[paragraph_dict['[OBJECTIVE]']]
    purpose.text = report.objective

    ############################################# AS IS PROCESS DESCRPTION: PROCESS OVERVIEW

    if report.process_overview:
        title= doc.paragraphs[paragraph_dict['[TITLE]']]
        title.text = execution.process_discovery.title
        title.style='Normal'
    else:
        doc.paragraphs[paragraph_dict['[TITLE]']].clear()

    ############################ AS IS PROCESS DESCRPTION: APPLICATIONS USED
    if report.applications_used:
        nameapps= doc.paragraphs[paragraph_dict['[DIFERENT NAMEAPPS]']]
        nameapps.text = "The applications used by the user during the execution of the process are:"
        applications_used(nameapps, execution, scenario, colnames)
        #nameapps.style = doc.styles['ListBullet'] --> add_paragraph('text', style='ListBullet')
    else:
        doc.paragraphs[paragraph_dict['[DIFERENT NAMEAPPS]']].clear()

    ##########################3 AS IS PROCESS MAP
    if report.as_is_process_map:
        bpmn= doc.paragraphs[paragraph_dict['[.BPMN]']]
        path_to_tree_file = os.path.join(execution.exp_folder_complete_path, scenario+"_results", "bpmn.dot")
        bpmn.text = f"Below is the process model in BPMN format associated with the case study with a total of {len(extract_all_activities_labels(path_to_tree_file))} activities and {len(extract_prev_act_labels(path_to_tree_file))} decision points."
        run = bpmn.add_run()
        run.add_picture(dot_to_png(path_to_tree_file), width=Inches(6))
        run.add_break()
    else:
        # Eliminar el párrafo [.BPMN] si no se cumple la condición
        doc.paragraphs[paragraph_dict['[.BPMN]']].clear()
        # O alternativamente eliminar el párrafo por completo
        #doc.paragraphs.pop(paragraph_dict['[.BPMN]'])

    
    #############################3 DETAILS AS IS PROCESS ACTIONS
    
    if report.detailed_as_is_process_actions:
        #meter diagrama de decision tree
        decision_tree= doc.paragraphs[paragraph_dict['[DECISION TREE]']]
        decision_tree.text = ""
        # path_to_tree_file = os.path.join(execution.exp_folder_complete_path, scenario+"_results", "decision_tree.pkl")
        # run = decision_tree.add_run()
        # run.add_picture(tree_to_png(path_to_tree_file), width=Inches(6))
        # run.add_break()
        #decision_tree.text = ''
        #detailes_as_is_process_actions(doc, paragraph_dict, execution, scenario)

        detailes_as_is_process_actions(doc, paragraph_dict, scenario, execution, colnames)
    else:
        doc.paragraphs[paragraph_dict['[DECISION TREE]']].clear()
     
    
    #############################3 INPUT DATA DESCRPTION
    if report.input_data_description:
        original_log= doc.paragraphs[paragraph_dict['[ORIGINAL LOG]']]
        df_logcsv = read_ui_log_as_dataframe(os.path.join(execution.exp_folder_complete_path, scenario, 'log.csv'))
        input_data_descrption(doc, original_log, execution, scenario, df_logcsv)
    else:
        doc.paragraphs[paragraph_dict['[ORIGINAL LOG]']].clear()


    doc.save(report_path)

    convert_docx_to_pdf(execution.user, report_path, report_path.replace('.docx', '.pdf'))

#################################################################################

def applications_used(nameapps, execution, scenario, colnames):
    logcsv_directory = os.path.join(execution.exp_folder_complete_path, scenario, 'log.csv')
    df_logcsv = read_ui_log_as_dataframe(logcsv_directory)
    unique_names = df_logcsv[colnames['NameApp']].unique()

    for name in unique_names:
        
        nameapps.add_run().add_break()
        nameapps.add_run('\n• ' + name, style='Título 4 Car')
        # Añadir salto de línea después del nombre
        nameapps.add_run().add_break()

        # Obtener la primera fila correspondiente al nombre de la aplicación
        row = df_logcsv[df_logcsv[colnames['NameApp']] == name].iloc[0]
        screenshot_filename = row[colnames['Screenshot']]
        screenshot_directory = os.path.join(execution.exp_folder_complete_path, scenario, screenshot_filename)

        # Añadir la imagen si existe
        if os.path.exists(screenshot_directory):
            nameapps.add_run().add_picture(screenshot_directory, width=Inches(6))
            # Añadir salto de línea después de la imagen
            nameapps.add_run().add_break()

        # Mensaje en caso de que no se encuentre la imagen
        else:
            print(f"Image not found for {name}: {screenshot_directory}")

#############################################################################################
def input_data_descrption(doc, original_log, execution, scenario, df_logcsv):
        
        table = doc.add_table(rows=(df_logcsv.shape[0] + 1), cols=df_logcsv.shape[1])

        # Insertar los nombres de las columnas
        for j, col in enumerate(df_logcsv.columns):
            table.cell(0, j).text = col

        # Insertar los datos del DataFrame
        for i in range(df_logcsv.shape[0]):
            for j in range(df_logcsv.shape[1]):
                table.cell(i + 1, j).text = str(df_logcsv.iloc[i, j])

        tbl, p = table._tbl, original_log._p
        p.addnext(tbl)
###################

def convert_docx_to_pdf(user, dx_path, pdf_path):
    try:
        if os.name == 'nt':
            convert(dx_path, pdf_path)
        else:
            subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), dx_path])
    except:
        create_notification(user, _('Limited report'), _('Please install MS Word on windows or Libreoffice on linux to generate the execution report.'), "")
        doc = aw.Document(dx_path)
        doc.save(pdf_path)


# def convert_docx_to_pdf2(dx_path, pdf_path):
#     extra_args = ['--pdf-engine-opt', '-dPDFSETTINGS=/prepress']
#     #pypandoc.convert_file(dx_path, 'pdf', outputfile=pdf_path, extra_args=extra_args)
#     pypandoc.convert_file(dx_path, 'pdf', outputfile=pdf_path, extra_args=extra_args)
##################################################

def detailes_as_is_process_actions(doc, paragraph_dict, scenario, execution, colnames):
    
    # def variant_column(df):
    #     sequence_to_variant = {}
    #     next_variant_number = 1
    #     # Función para generar el mapeo de variantes y asignar valores a la columna Variant2
    #     def assign_variant(trace):
    #         nonlocal next_variant_number  # Declarar next_variant_number como global
    #         cadena = ""
    #         for e in trace[colnames['Activity']].tolist():
    #             cadena = cadena + str(e)
    #         if cadena not in sequence_to_variant:
    #             sequence_to_variant[cadena] = next_variant_number
    #             next_variant_number += 1
    #         trace[colnames['Variant']] = sequence_to_variant[cadena]
    #         return trace

    #     # Aplicar la función a cada grupo de trace_id y asignar el resultado al DataFrame original
    #     df=df.groupby(colnames['Case']).apply(assign_variant).reset_index(drop=True)
    #     return df
    
    def cambiar_color_nodos_y_caminos(labels, path_to_dot_file):
        # Cargar el contenido del archivo .dot
        with open(path_to_dot_file, 'r') as file:
            dot_content = file.read()
        # Crear un grafo desde el contenido del archivo .dot
        grafo = pgv.AGraph(string=dot_content)
        
        # Crear un conjunto de nodos que necesitamos colorear
        nodos_a_colorear = set()
        labels = [str(elemento) for elemento in labels]
        # Recorrer todos los nodos del grafo y colorear los nodos correspondientes
        for nodo in grafo.nodes():
            if nodo.attr['label'] in labels:
                nodo.attr['fillcolor'] = "yellow"
                nodo.attr['style'] = 'filled'
                nodos_a_colorear.add(nodo.name)
            elif nodo.attr['label'] == "":
                nodo.attr['label'] = " "
                nodo.attr['fillcolor'] = nodo.attr['fillcolor']
                nodo.attr['style'] = nodo.attr['style']
            else:
                nodo.attr['style'] = 'filled'
                

        # Recorrer todas las aristas del grafo y colorear las que conectan nodos en la lista
        # Identificar y colorear los nodos de tipo "diamond" entre los nodos etiquetados
        for edge in grafo.edges():
            if edge[0] in nodos_a_colorear and edge[1] not in nodos_a_colorear:
                nodo_destino = grafo.get_node(edge[1])
                if nodo_destino.attr['shape'] == 'diamond':
                    nodo_destino.attr['fillcolor'] = "yellow"
                    nodo_destino.attr['style'] = 'filled'
                    nodos_a_colorear.add(nodo_destino.name)
            elif edge[1] in nodos_a_colorear and edge[0] not in nodos_a_colorear:
                nodo_origen = grafo.get_node(edge[0])
                if nodo_origen.attr['shape'] == 'diamond':
                    nodo_origen.attr['fillcolor'] = "yellow"
                    nodo_origen.attr['style'] = 'filled'
                    nodos_a_colorear.add(nodo_origen.name)

        # Colorear las aristas conectadas entre los nodos coloreados
        for edge in grafo.edges():
            if edge[0] in nodos_a_colorear and edge[1] in nodos_a_colorear:
                edge.attr['color'] = "blue"
                edge.attr['penwidth'] = 2.0

        # Configurar atributos globales del grafo
        grafo.graph_attr.update(bgcolor='white', rankdir='LR')
        grafo.graph_attr['overlap'] = 'false'
        grafo.format = 'png'

        # Guardar la imagen a un archivo temporal
        temp_file = NamedTemporaryFile(delete=False, suffix='.png')
        grafo.draw(temp_file.name, prog='dot', format='png')
        graph_path = temp_file.name

        return graph_path

    def get_branch_condition(decision_point, branch_number):
        branches = decision_point.get('branches', [])
        rules = decision_point.get('rules', {})
        
        for branch in branches:
            if branch['label'] == str(branch_number):
                res= rules.get(str(branch_number))
                break
            else: 
                res= "No se han encontrado las reglas asociadas a esta rama."
        return res
##el resultado es un diccionario cuyas claves son las reglas y los valores otro diccionario con las reglas y su valor va a ser el centroid
##{'numeric__Coor_Y_2 > 382.39 & numeric__Coor_Y_3 > 491.51': {2: [['numeric__Coor_Y_2 > 382.39', 'aqui va el centroid en tupla']], 3: [['numeric__Coor_Y_3 > 491.51', 'aqui va el centroid en tupla']]}}
    def get_branch_condition2(rules, pre_pd_activities, process):
        def coordinate_rule(variable, condition, condition_dict, exists):
            # Verificar si el final de la variable contiene una de las actividades finales con o sin sufijo adicional
            present = False
            for act in pre_pd_activities:
                pattern = re.compile(f"_{act}(?:_.*)?$")
                if pattern.search(variable):
                    variable_parts = variable.split('-')
                    
                    if len(variable_parts) > 1:
                        last_part = variable_parts[-1]
                        
                        first_part_split = variable_parts[-2].split('_')
                        last_part_split = last_part.split('_')
                        
                        first_element = first_part_split[-1]
                        second_element = last_part_split[0]

                        compo_class_or_text = "_".join(last_part_split[2:]) # last_part_split[1] is the activity for the component
                        if compo_class_or_text == "NaN" or compo_class_or_text == "nan":
                            exists = not exists # Invert the exists value if the component is NaN, meaning <= is actually compo exists and > not exists
                        
                        if act not in condition_dict:
                            condition_dict[act] = []
                        condition_dict[act].append(["centroid", condition, (first_element, second_element, compo_class_or_text), exists])
                        present = True

        result_dict = {}
        rules = preprocess_rules(rules, process)
        for rule in rules:
            condition_dict = {}
            parts = rule.split('&')
            for part in parts:
                elements = re.split(r'<=|>=|<|>|==|!=', part)
                variable = elements[0].strip()
                condition = part.strip()

                # Figure out weather the element is or not in this instance
                exists = False
                if re.search(r">=|>|==", part):
                    exists = True

                # We technically should have no dp rules left at this point
                pattern_with_centroid = r"([a-zA-Z_]+)__([a-zA-Z0-9_-]+)_(\d+\.?\d*?-\d+\.?\d*?)_(\d+)(_?[_0-9a-zA-Z]+)"
                if re.match(pattern_with_centroid, variable):
                    coordinate_rule(variable, condition, condition_dict, exists)
            result_dict[rule.strip()] = condition_dict
        return result_dict


#hay que decidr que en el punto de decisión hay una regla solpada entre dos branchas, x e x, y explicar esas reglas.
#le paso una rama y me da un diccioanrio con las reglas solapadas que tiene con cada una de las otras ramas del punto de decision
#hay que sacar un diccionario con esta forma: {'6':{'numeric__Coor_Y_2 > 382.39 & numeric__Coor_Y_3 > 491.51': {2: [['numeric__Coor_Y_2 > 382.39', 'numeric__Coor_Y_2']], 3: [['numeric__Coor_Y_3 > 491.51', 'numeric__Coor_Y_3']]}}}
    def get_overlapping_branch_condition2(decision_point, branch_number, pre_pd_activities):
        branches = decision_point.get('branches', [])
        rules = decision_point.get('overlapping_rules', {})
        branch_rules = defaultdict(list)
        result_dict = defaultdict(dict)
        
        # Extract the rules for the given branch number
        target_branch_rules = rules.get(str(branch_number), [])
        if not target_branch_rules:
            return {"No associated rule found": []}
        
        # Parse rules for the given branch number and collect conditions
        for rule in target_branch_rules:
            condition_dict = {}
            parts = rule.split('&')
            for part in parts:
                elements = re.split(r'<=|>=|<|>|==|!=', part)
                variable = elements[0].strip()
                condition = part.strip()
                
                matched = False
                for act in pre_pd_activities:
                    pattern = re.compile(f"_{act}(?:_.*)?$")
                    if pattern.search(variable):
                        matched = True
                        variable_parts = variable.split('-')
                        
                        if len(variable_parts) > 1:
                            last_part = variable_parts[-1]
                            
                            first_part_split = variable_parts[0].split('_')
                            last_part_split = last_part.split('_')
                            
                            first_element = first_part_split[-1]
                            second_element = last_part_split[0]
                            
                            if act not in condition_dict:
                                condition_dict[act] = []
                            condition_dict[act].append([condition, (first_element, second_element)])
                if not matched:
                    pass
            branch_rules[rule] = condition_dict
        
        # Find and map overlapping rules with other branches
        for other_branch, other_branch_rules in rules.items():
            if other_branch != str(branch_number):
                for rule in other_branch_rules:
                    if rule in target_branch_rules:
                        result_dict[other_branch][rule] = branch_rules[rule]
        
        return dict(result_dict)
    
    
    def draw_polygons_on_image(json_data, image_path, doc_decision_tree):
        # Cargar la imagen
        image = Image.open(image_path)
        print(json_data)
        # Dibujar los polígonos en la imagen
        draw = ImageDraw.Draw(image)
        for e in json_data:
            points = [tuple(point) for point in e["points"]]
            color = e.get("color", (0, 255, 0)) #el color verde por defecto
            if len(points) > 2:
                draw.polygon(points, outline=color, width=4)
            else:
                draw.circle(points[0], radius=5 ,fill=color)

        # Guardar la imagen modificada en memoria
        image_byte_array = io.BytesIO()
        image.save(image_byte_array, format='PNG')
        image_byte_array.seek(0)

        # Añadir la imagen al documento
        doc_decision_tree.add_run().add_break()
        doc_decision_tree.add_run().add_picture(image_byte_array, width=Inches(6))
        doc_decision_tree.add_run().add_break()
    
## devuelve un diccionario cuyas claves son las prev act delos punto de decisiones que hay en una varianye y 
# de valor los json del punto de dceision

    def get_decision_points_for_branches(data, branch_labels):
        result = {}

        def search_decision_points(decision_points, branch_labels):
            for dp in decision_points:
                for branch in dp['branches']:
                    if str(branch['label']) in branch_labels:
                        result[dp['prevAct']] = dp
                    search_decision_points(branch.get('decision_points', []), branch_labels)

        search_decision_points(data['decision_points'], branch_labels)
        return result
    
    
#extrae una lista con todos los id de todos los puntos de decision de una variante  
    def extract_decision_point_ids(decision_points):
        ids = []
    
        def extract_ids(dp_list):
            for dp in dp_list:
                ids.append(dp['id'])
                for branch in dp['branches']:
                    extract_ids(branch['decision_points'])
        
        extract_ids(decision_points)
        return ids
    #def explicabilidad_decisions(decision_points, variant_decision_points):


#extrae una lista con todos los id de todos los puntos de decision de una variante  
    def extract_decision_point_ids(decision_points):
        ids = []
    
        def extract_ids(dp_list):
            for dp in dp_list:
                ids.append(dp['id'])
                for branch in dp['branches']:
                    extract_ids(branch['decision_points'])
        
        extract_ids(decision_points)
        return ids
    #def explicabilidad_decisions(decision_points, variant_decision_points):


    def explicabilidad_actions(decision_tree, activity,group, colnames):

        def pintar_imagen(k, action_dict):
            event_description = None
            image_click = None
            out_click=None
            image = None

            action = pd.DataFrame(action_dict[k])
            #action= action_dict[k]

            decision_tree.add_run().add_break()
            decision_tree.add_run(f'Action {k}', style='Título 5 Car')
            decision_tree.add_run().add_break()
            
            if action[colnames['EventType']].iloc[0] == 1 or action[colnames['EventType']].iloc[0] == "click": #colnames['EventType']
                # Calcular la media de Coor_X y Coor_Y
                mean_x = action[colnames['CoorX']].mean()
                mean_y = action[colnames['CoorY']].mean()

                # Cargar la imagen correspondiente
                screenshot_filename = action[colnames['Screenshot']].iloc[0]
                path_to_image = os.path.join(execution.exp_folder_complete_path, scenario, screenshot_filename)
                image = action[colnames['Screenshot']].iloc[0]
                image_click=True
                with Image.open(path_to_image) as img:
                    width, height = img.size
                    print(width, height)
                if width>=mean_x and height>=mean_y:
                    event_description = f"The user clicks at point ({mean_x}, {mean_y})"
                else:
                    if mean_x > width:  side="right"# Click is to the right of the image       
                    if mean_y > height:  side="bottom"# Click is below the image      
                    if mean_x < 0:  side="left"# Click is to the left of the image         
                    if mean_y < 0: side="top"# Click is above the image

                    event_description = f"ERROR: coordinates recorded incorrectly, out of screen resolution by {side} border (highlighted in red in the picture). User clicks on {mean_x}, {mean_y} and the screen resolution is {width}, {height}."
                    out_click=True

            else:
                    # Mostrar el valor de TextInput si existe, de lo contrario imprimir "No TextInput"
                text_input_cat_name = execution.ui_elements_classification.model.text_classname
                text_input = action[text_input_cat_name].iloc[0] if text_input_cat_name in action.columns and not pd.isnull(action[text_input_cat_name].iloc[0]) else f"No {text_input_cat_name}"
                event_description = f'The user writes "{text_input}"'
                screenshot_filename = action[colnames['Screenshot']].iloc[0]
                path_to_image = os.path.join(execution.exp_folder_complete_path, scenario, screenshot_filename)
                image = action[colnames['Screenshot']].iloc[0] if 'Screenshot' in action.columns else None
                

            
            decision_tree.add_run().add_break()
            if event_description:
                decision_tree.add_run(event_description + '\n')
            decision_tree.add_run().add_break()

            if image:
                with Image.open(path_to_image) as img:
                    draw = ImageDraw.Draw(img)
                    
                    if image_click and not out_click: # si en la imagen se encuentra el click
                        
                        # Dibujar un pequeño cuadrado alrededor de las coordenadas medias
                        box_size = 10  # Ajustar el tamaño del cuadrado según sea necesario
                        left = mean_x - box_size / 2
                        top = mean_y - box_size / 2
                        right = mean_x + box_size / 2
                        bottom = mean_y + box_size / 2
                        
                        draw.rectangle([left, top, right, bottom], outline="red", width=5)
                        # Convertir la imagen a un objeto byte para insertar en docx   
                    if out_click: # If the click is outside the image
                        border_width = 10
                        if side=="right":  # Click is to the right of the image
                            draw.line([(width - 1, 0), (width - 1, height)], fill="red", width=border_width)
                        if side=="bottom":  # Click is below the image
                            draw.line([(0, height - 1), (width, height - 1)], fill="red", width=border_width)
                        if side=="left":  # Click is to the left of the image
                            draw.line([(0, 0), (0, height)], fill="red", width=border_width)
                        if side=="top":  # Click is above the image
                            draw.line([(0, 0), (width, 0)], fill="red", width=border_width)

                image_stream = io.BytesIO()
                img.save(image_stream, 'PNG')
                image_stream.seek(0)
                decision_tree.add_run().add_picture(image_stream, width=Inches(6))
                decision_tree.add_run().add_break()
        #####################################

        decision_tree.add_run().add_break()
        decision_tree.add_run(f'Activity {activity}\n', style='Título 4 Car')
        
        decision_tree.add_run().add_break()

        ############################ EXPLICABILIDAD DE LAS ACTIVIDADES Y ACCIONES
        # Filtrar por actividad
        activity_group = group[group[colnames['Activity']] == activity]

        # Agrupar por trace_id dentro de activity_group
        activity_actions_group = activity_group.groupby(colnames['Case'])
        
        # Verificar si todos los grupos tienen una fila
        all_single_action = all(len(actions) == 1 for _, actions in activity_actions_group)
        # Iterar sobre cada grupo y determinar el número de acciones

        

        if all_single_action: # ACTIVIDAD CON UNA ACCION
            action_dict = {}
            for i, (activity_label, action) in enumerate(activity_group.groupby(colnames['Activity']), start=1):
                action_dict[i] = action

            for k in sorted(action_dict.keys()):
                pintar_imagen(k, action_dict)    
                
             
        else: #ACTIVIDADES CON MAS DE UNA ACCION 
            action_dict = {}  # Crear el diccionario vacío    
            for i, (activity_label, action_group) in enumerate(activity_actions_group):
                #crear un diccionario, cuyos claves sean el indice de accion (accion 1, 2, 3) y los valores sea el grupo de filas de la accion
                            
                for j, (index, action) in enumerate(action_group.iterrows(), start=1):
                    #relleno el diccionario
                    
                    if j not in action_dict:
                        action_dict[j] = []
                    # Relleno el diccionario
                    action_dict[j].append(action)
                    #action_dict[j] = action_dict[j].append(action[1])

            #recorrer las claves del diccionario
            for k in sorted(action_dict.keys()):
                pintar_imagen(k, action_dict)
          
############################################################3
    
    def process_variant_group(group, traceability, path_to_dot_file, colnames,df_pd_log, variant):
        def process_centroid_conditions(centroid_conditions, activities):
            for act, condition_list in centroid_conditions.items(): 
                color_index = 0
                compo_ui_json=[]
                screenshot_filename=None
                #obtenemos la screenshot correspondiente
                ######################################3
                if act in activities:
                    screenshot_filename = group[group[colnames['Activity']] == act][colnames['Screenshot']].iloc[0]
                else:
                    screenshot_filename = df_pd_log[df_pd_log[colnames['Activity']] == act][colnames['Screenshot']].iloc[0]
                if not screenshot_filename:
                    print(f"No screenshot found for the activity '{act}'.")
                    ##################3
                path_to_image = os.path.join(execution.exp_folder_complete_path, scenario, screenshot_filename)
                for _, condition, variable, existence in condition_list:
                    #decision_tree.add_run(f'Activity {act}: {condition}\n')
                    color = colors[color_index % len(colors)]
                    # Insertar la imagen en el documento
                    ui_compo_centroid=variable[:2]
                    ui_compo_class_or_text=variable[2]
                    # Correct the text if it is NaN
                    aux_compo_text = "Some component" if ui_compo_class_or_text == "NaN" else ui_compo_class_or_text
                    
                    if existence:
                        compo_ui = get_uicompo_from_centroid(screenshot_filename, ui_compo_centroid, ui_compo_class_or_text, os.path.join(execution.exp_folder_complete_path, scenario + '_results'))
                        if compo_ui:
                            compo_ui["color"] = color  # Asignar color al polígono
                            compo_ui_json.append(compo_ui)

                            run = decision_tree.add_run(f'Existence of "{aux_compo_text}" with centroid {ui_compo_centroid} in activity {act}\n')
                            run.font.color.rgb = RGBColor(*color)
                            color_index += 1
                            continue
                    # Fallback if not exists or component not found
                    compo_ui = {"points": [[int(float(coord)) for coord in ui_compo_centroid]], "color": color}
                    compo_ui_json.append(compo_ui)

                    run = decision_tree.add_run(f'Non-existence of "{aux_compo_text}" with centroid {ui_compo_centroid} in activity {act}\n')
                    run.font.color.rgb = RGBColor(*color)
                    color_index += 1

                if compo_ui_json:
                    draw_polygons_on_image(compo_ui_json, path_to_image,decision_tree)
                    decision_tree.add_run().add_break()


        def process_decision_conditions(decision_conditions, activities):
            for branch, existence in decision_conditions:
                color = colors[0]
                if existence:
                    run = decision_tree.add_run(f'Took the branch "{branch}" in a previous decision point\n')
                else:
                    run = decision_tree.add_run(f'Did not take the branch "{branch}" in a previous decision point\n')
                run.font.color.rgb = RGBColor(*color)

        activities = group[colnames['Activity']].unique().tolist() #en funcion de si las activity label se pueden repetir

        variant_decision_points = group.filter(regex=r'id[a-zA-Z0-9]{8}-[a-zA-Z0-9]{4}-[a-zA-Z0-9]{4}-[a-zA-Z0-9]{4}-[a-zA-Z0-9]+').dropna(axis=1, how='all')

        dps = list(filter(lambda dp: dp.id in variant_decision_points.columns, traceability.get_non_empty_dp_flattened()))

        #prev_act = traceability['decision_points'][0]['prevAct']
        # decision_point = traceability['decision_points'][0]
        
        # #print(decision_point)
        # first_dp_id= decision_point['id']
        # #ir acumulando las id de los puntos de decision de cada variante
        # columnas_label_ramas=[]
        # for pd in list(map(lambda x: x['id'], traceability['decision_points'])):
        #     columnas_label_ramas= columnas_label_ramas + group[pd].unique().tolist()
        #     variant_decision_points = get_decision_points_for_branches(traceability, columnas_label_ramas)
        # #añadir el primer punto de decision al diccionario
        # variant_decision_points[decision_point['prevAct']]= decision_point
        
        decision_tree.add_run().add_break()
        decision_tree.add_run(f'Variant {variant}\n', style='Título 3 Car')
        decision_tree.add_run().add_break()

        #actividades
        
        #meto diagrama bpmn resaltado de variantes
        run = decision_tree.add_run()
        run.add_picture(cambiar_color_nodos_y_caminos(activities, path_to_dot_file), width=Inches(6))
        run.add_break()

        for i, activity in enumerate(activities):

            #EXPLICAR ACTIVIDADES Y ACCIONES
            explicabilidad_actions(decision_tree, activity,group, colnames)
            
            ###############################################################################################333

        #EXPLICABILIDAD DE LAS DECISIONES
        for dp in dps:
            activity = str(dp.prevAct)
            num_ramas = len(dp.branches)
            target = str(variant_decision_points[dp.id].iloc[0])
            branch_rules = list(filter(lambda r: str(r.target) == target, dp.rules))[0].condition
            #{'numeric__Coor_Y_2 <= 382.39': {2: [['numeric__Coor_Y_2 <= 382.39', 'numeric__Coor_Y_2']]}}
            #actividades previas al pd, independientemente de la variante
            act_seq_index = activities.index(int(activity))
            pre_activities = activities[:act_seq_index+1]

            result_dict = get_branch_condition2(branch_rules, pre_activities, traceability)
            result_dict_overlapping = {}
            # result_dict_overlapping = get_overlapping_branch_condition2(decision_point, next_activity, pre_activities)

            decision_tree.add_run().add_break()      
            decision_tree.add_run(f'Decision Point\n', style='Título 4 Car').bold = True
            decision_tree.add_run().add_break()
            decision_tree.add_run().add_break()      
            decision_tree.add_run(f'In this variant, after activity {activity} there is a decision point where the user decides between taking {num_ramas} different branches. \n')
            decision_tree.add_run(f'In the case of this variant (variant {variant}) the user chooses to take branch {target}. \n')
            decision_tree.add_run().add_break()
            if 'No associated rule found' not in result_dict:
                decision_tree.add_run(f"In order for the user to decide this branch, one of these deterministic rules must be given:")
                for rules in result_dict.keys():
                    decision_tree.add_run('\n• ' + f"{rules}" + '\n')    
                decision_tree.add_run().add_break()
            else:
                decision_tree.add_run(f"No associated deterministic rules are found:")
                decision_tree.add_run().add_break()
            if 'No associated rule found' not in result_dict_overlapping:
                for branch, rules in result_dict_overlapping.items(): 
                    decision_tree.add_run(f"In this decision we find rules that overlap with the rest of the decision branches. These rules are:")
                    for rule, conditions in rules.items():
                        decision_tree.add_run('\n• ' + f"Rule overlapping with branch {branch}: {rule}." + '\n')
                    decision_tree.add_run().add_break()
            else:
                decision_tree.add_run(f"No overlapping rules were found.")
                decision_tree.add_run().add_break()


            colors = [(255, 0, 0), (0, 255, 0), (0, 0, 255), (255, 165, 0)]

            #################################################################################################### DETERMINIST RULES
            
            if result_dict and all(isinstance(value, dict) for value in result_dict.values()): #se ejecuta solo si tiene reglas asociadas
                decision_tree.add_run().add_break()
                decision_tree.add_run("Determinist Rules\n", style='Título 5 Car').bold = True
                decision_tree.add_run().add_break()
                print(result_dict)
                for rule, conditions in result_dict.items():
                    decision_tree.add_run(f'{rule}\n').bold = True
                    #compo_ui_json=[]
                    #color_index = 0
                    centroid_conditions = {k: v for k, v in conditions.items() if v[0][0] == "centroid"}
                    decision_conditions = conditions["decision"] if "decision" in conditions else []
                    process_centroid_conditions(centroid_conditions, activities)
                    process_decision_conditions(decision_conditions, activities)
                decision_tree.add_run().add_break()

            #################################################################################################### OVERLAPPED RULES

            if 'No associated rule found' not in result_dict_overlapping:  # se ejecuta solo si tiene reglas asociadas
                decision_tree.add_run().add_break()
                decision_tree.add_run("Overlapping Rules\n", style='Título 5 Car').bold = True
                decision_tree.add_run().add_break()
                print(result_dict_overlapping)
                for branch, rules in result_dict_overlapping.items():
                    decision_tree.add_run(f'Branch {branch}:\n').bold =True
                    
                    for rule, conditions in rules.items():
                        decision_tree.add_run(f'{rule}\n').bold = True
                        color_index = 0
                        compo_ui_json = []
                        screenshot_filename= None
                        for act, condition_list in conditions.items():
                            # Obtener la screenshot correspondiente
                            if act in activities:
                                screenshot_filename = group[group[colnames['Activity']] == act][colnames['Screenshot']].iloc[0]
                            else:

                                screenshot_filename = df_pd_log[df_pd_log[colnames['Activity']] == act][colnames['Screenshot']].iloc[0]
                            if not screenshot_filename:
                                print(f"No screenshot found for the activity '{act}'.")

                            path_to_image = os.path.join(execution.exp_folder_complete_path, scenario, screenshot_filename)

                            #TODO: Adapt to new format of condition list after fixing overlapping rules
                            for condition, variable in condition_list:
                                color = colors[color_index % len(colors)]
                                run = decision_tree.add_run(f'Activity {act}: {condition}\n')
                                run.font.color.rgb = RGBColor(*color)
                                color_index += 1

                                # Insertar la imagen en el documento
                                ui_compo_centroid = variable
                                compo_ui = get_uicompo_from_centroid(screenshot_filename, ui_compo_centroid, os.path.join(execution.exp_folder_complete_path, scenario + '_results'))

                                if compo_ui:
                                    compo_ui["color"] = color  # Asignar color al polígono
                                    compo_ui_json.append(compo_ui)

                            if compo_ui_json:
                                draw_polygons_on_image(compo_ui_json, path_to_image, decision_tree)
                                decision_tree.add_run().add_break()

                        decision_tree.add_run().add_break()
    
###############
    
    decision_tree= doc.paragraphs[paragraph_dict['[DECISION TREE]']]
    decision_tree.text = ''
    df_pd_log = read_ui_log_as_dataframe(os.path.join(execution.exp_folder_complete_path, scenario+'_results', PROCESS_DISCOVERY_LOG_FILENAME))
    #se quita porque la columna se aplica ya cuando se crea el csv y se llama auto_variant
    #df2= variant_column(df)
    traceability_json = json.load(open(os.path.join(execution.exp_folder_complete_path, scenario+'_results', 'traceability.json')))
    traceability = Process.from_json(traceability_json)
    path_to_dot_file = os.path.join(execution.exp_folder_complete_path, scenario+"_results", "bpmn.dot")
    #df2.groupby('Variant2').apply(lambda group: process_variant_group(group,traceability,path_to_dot_file, colnames))
    #colnames['Variant'] es auto_variant
    df_pd_log.groupby(colnames['Variant'], as_index=False).apply(lambda group: process_variant_group(group,traceability,path_to_dot_file, colnames,df_pd_log, group.name))


###################################################################
#####################################################################

def deleteReport(request):
    report_id = request.GET.get("report_id")

    removed_report = PDD.objects.get(id=report_id)
    if request.user.id != removed_report.user.id:
        raise Exception(_("This case study doesn't belong to the authenticated user"))
    
    #if removed_report.executed != 0:
    #    raise Exception(_("This case study cannot be deleted because it has already been excecuted"))

    execution = removed_report.execution

    scenarios_to_study = execution.scenarios_to_study


    for scenario in scenarios_to_study:
        report_directory = os.path.join(execution.exp_folder_complete_path, scenario+"_results")
        report_path = os.path.join(report_directory, f'report_{report_id}.docx')
        if os.path.exists(report_path):
            os.remove(report_path)
        report_path_pdf = os.path.join(report_directory, f'report_{report_id}.pdf')
        if os.path.exists(report_path_pdf):
            os.remove(report_path_pdf)

    removed_report.delete()

    return HttpResponseRedirect(reverse("analyzer:execution_detail", kwargs={"execution_id": removed_report.execution.id}))

#############################################################################################3



class ReportingConfigurationDetail(DetailView):
    def get(self, request, *args, **kwargs):
        report = get_object_or_404(PDD, id=kwargs["report_id"])
        
        form = ReportingForm(read_only=True, instance=report)  
        context = {"form": form,
            "execution": report.execution,
            }
        return render(request, 'reporting/detail.html', context)
    

##################################################3

def download_report_zip(request, report_id):
    report = get_object_or_404(PDD, pk=report_id)
    execution= report.execution

    zip_filename = f"execution_{execution.id}_reports_{report.id}.zip"

    zip_path = os.path.join(settings.MEDIA_ROOT, zip_filename)

    with zipfile.ZipFile(zip_path, 'w') as zipf:
        scenarios_to_study = execution.scenarios_to_study

        for scenario in scenarios_to_study:

            report_directory = os.path.join(execution.exp_folder_complete_path, scenario + "_results")
            report_filename = f'report_{report.id}.docx'
            report_path = os.path.join(report_directory, report_filename)
            
            new_filename = f'scenario_{scenario}_report_{report.id}.docx'

            # Ensure the file exists before adding to ZIP
            if os.path.exists(report_path):
                zipf.write(report_path, arcname=new_filename)

    with open(zip_path, 'rb') as f:
        response = HttpResponse(f.read(), content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename={zip_filename}'

    os.remove(zip_path)  # Clean up the generated zip file after serving it
    return response


##################################################
def preview_pdf(request, report_id):
    
    report = get_object_or_404(PDD, pk=report_id)
    execution = report.execution
    pdf_path = os.path.join(execution.exp_folder_complete_path, execution.scenarios_to_study[0]+'_results', f'report_{report.id}.pdf')
    #pdf_path = '/screenrpa/media/unzipped/datos_inciiales_j49gvQs_1714120837/executions/exec_41/sc_0_size50_Balanced_results/calendario-academico.pdf'

    return FileResponse(open(pdf_path, 'rb'), content_type='application/pdf')

    # if os.path.exists(pdf_path):
    #     # Renderizar una plantilla que contenga el iframe
    #     return render(request, 'reporting/preview_pdf.html', {'pdf_path': pdf_path})
    # else:
    #     return HttpResponseNotFound("El documento PDF solicitado no fue encontrado.")

def preprocess_rules(rules, process):
    def unpack_decision_rule(variable, condition, rules, exists, process):
        coincidences = re.match(r"([a-zA-Z_]+)__([a-zA-Z0-9-]+)_([a-zA-Z]+)?_?([_a-zA-Z0-9-]+)", variable)
        target = coincidences.group(4) # Actividad o numero de puerta xor
        if coincidences.group(3):  # Si hay un grupo 3 adicional (opcional, significa xor)
            target = f"{coincidences.group(3)}_{target}"
        
        dp = coincidences.group(2)
        dps = process.get_non_empty_dp_flattened()
        dp = list(filter(lambda x: x.id == dp, dps))[0]

        subrules = dp.to_json().get("rules", {})[target]
        # In case we have branch rules, we process them
        aux_rules = copy.deepcopy(rules)
        for subrule in subrules:
            for i, r in enumerate(aux_rules):
                # If we went to that branch we add the rules as &, else we add it as | and flip the sign
                # In our model it means adding a separate rule
                # Ej. numeric__Coor_Y_2 > 382.39 & numeric__Coor_Y_3 > 491.51 is the branch condition
                # Ej. numeric__Coor_Y_2 <= 382.39 | numeric__Coor_Y_3 <= 491.51 is the "not" branch condition
                if exists:
                    # We prevent duplicate rules
                    aux = [sr.strip() for sr in subrule.split('&') if sr not in r]
                    aux = " & ".join(aux)
                    # Replace old rule by new one. We only moddify rules that contain the variable
                    rules[i] = r.replace(condition, aux)
                else:
                    aux = [sr.strip() for sr in subrule.split('&') if sr not in r]
                    for j, sc in enumerate(aux):
                        if "<=" in sc:
                            aux[j] = sc.replace("<=", ">")
                        elif ">=" in sc:
                            aux[j] = sc.replace(">=", "<")
                        elif "<" in sc:
                            aux[j] = sc.replace("<", ">=")
                        elif ">" in sc:
                            aux[j] = sc.replace(">", "<=")
                        elif "==" in sc:
                            aux[j] = sc.replace("==", "!=")
                        elif "!=" in sc:
                            aux[j] = sc.replace("!=", "==")
                        # To avoid messing up with indexes we moddify the first, add the rest at the end
                        if j == 0:
                            rules[i] = r.replace(condition, aux[j])
                        else:
                            rules.append(r.replace(condition, aux[j]))
        # Before returning, if case new branch rules were added, we need to preprocess them
        rules = preprocess_rules(rules, process)

    # To unpack dp rules and treat everything correctly we need to first trat all dp rules then the rest
    aux_rules = copy.deepcopy(rules)
    for rule in aux_rules:
        parts = rule.split('&')
        for part in parts:
            elements = re.split(r'<=|>=|<|>|==|!=', part)
            variable = elements[0].strip()
            condition = part.strip()

            # Figure out weather the element is or not in this instance
            exists = False
            if re.search(r">=|>|==", part):
                exists = True
            
            pattern_with_centroid = r"([a-zA-Z_]+)__([a-zA-Z0-9_-]+)_(\d+\.?\d*?-\d+\.?\d*?)_(\d+)(_?[_0-9a-zA-Z]+)"
            # numeric__id6322e007-a58b-4b5a-b711-8f51d37c438f_1
            pattern_decision_point = r"([a-zA-Z_]+)__([a-zA-Z0-9-]+)_([a-zA-Z]+)?_?([_a-zA-Z0-9-]+)"

            # Centroids will also match the decision point pattern, so we need to treat them first
            if re.match(pattern_with_centroid, variable):
                continue
            elif re.match(pattern_decision_point, variable):
                unpack_decision_rule(variable, condition, rules, exists, process)
    return rules